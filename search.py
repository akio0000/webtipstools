import json
from pathlib import Path
from datetime import datetime
import streamlit as st
import base64
import pandas as pd
import io
from docx import Document
from docx.shared import Pt, Inches, Cm
import re
from fpdf import FPDF
import zipfile
from utils import collect_existing_tags, collect_existing_participants, compress_video, get_db_connection, delete_record_from_db, update_record_in_db
from config import DEFAULT_DATA_DIR, DB_PATH

# ==========================================
# JSON 読み込み（フォルダ内の全 JSON）
# ==========================================
# ==========================================
# 検索処理 (DB版)
# ==========================================
def search_entries(conditions, mode="or"):
    conn = get_db_connection()
    try:
        query = "SELECT * FROM records WHERE 1=1"
        params = []
        
        sub_queries = []
        if conditions.get("title"):
            sub_queries.append("title LIKE ?")
            params.append(f"%{conditions['title']}%")
        
        if conditions.get("text"):
            sub_queries.append("content LIKE ?")
            params.append(f"%{conditions['text']}%")
            
        if conditions.get("tag"):
            sub_queries.append("tags LIKE ?")
            params.append(f"%{conditions['tag']}%")
            
        if conditions.get("source") and conditions.get("source") != "すべて":
            query += " AND source = ?"
            params.append(conditions["source"])

        if sub_queries:
            op = " OR " if mode == "or" else " AND "
            query += " AND (" + op.join(sub_queries) + ")"
        
        query += " ORDER BY created_at DESC"
        
        rows = conn.execute(query, params).fetchall()
        results = []
        for row in rows:
            entry = dict(row)
            # JSON形式のフィールドをデコード
            entry["text"] = entry.pop("content")
            entry["tags"] = entry["tags"].split(",") if entry["tags"] else []
            entry["file_path"] = json.loads(entry["file_path"])
            entry["details"] = json.loads(entry["details"])
            results.append(entry)
        return results
    finally:
        conn.close()

def delete_entry(record_id):
    return delete_record_from_db(record_id)

def update_entry(record_id, updated_fields):
    return update_record_in_db(record_id, updated_fields)

# ==========================================
# ファイルプレビュー表示
# ==========================================
def show_file_preview(raw_path, unique_key):
    if not raw_path:
        return

    file_path = raw_path.replace("\\", "/")
    file_path_obj = Path(file_path)
    file_name = file_path_obj.name

    if not file_path_obj.exists():
        st.warning(f"添付ファイルが見つかりません: {file_name}")
        return

    # 画像プレビュー
    if file_path.lower().endswith((".png", ".jpg", ".jpeg", ".gif")) and file_path_obj.is_file():
        with st.expander("画像を表示"):
            st.image(file_path, caption=file_name)
        return

    # 動画プレビュー
    if file_path.lower().endswith((".mp4", ".mov", ".avi", ".mkv")) and file_path_obj.is_file():
        with st.expander("動画を再生"):
            st.video(file_path)
        return

    # PDF プレビュー
    if file_path.lower().endswith(".pdf") and file_path_obj.is_file():
        with st.expander("PDFを表示"):
            try:
                with open(file_path_obj, "rb") as f:
                    base64_pdf = base64.b64encode(f.read()).decode('utf-8')
                pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600"></iframe>'
                st.markdown(pdf_display, unsafe_allow_html=True)
            except Exception as e:
                st.error(f"PDF読み込みエラー: {e}")
        return

    # CSV プレビュー
    if file_path.lower().endswith(".csv") and file_path_obj.is_file():
        with st.expander("CSVを表示"):
            try:
                df = pd.read_csv(file_path_obj, encoding="utf-8")
                st.dataframe(df)
            except Exception as e:
                st.warning(f"CSV読み込みエラー: {e}")
        return

    # Excel プレビュー
    if file_path.lower().endswith(".xlsx") and file_path_obj.is_file():
        with st.expander("Excelを表示"):
            try:
                df = pd.read_excel(file_path_obj)
                st.dataframe(df)
            except Exception as e:
                st.warning(f"Excel読み込みエラー: {e}")
        return

    # URL ショートカット (.url) プレビュー
    if file_path.lower().endswith(".url") and file_path_obj.is_file():
        with st.expander("URLショートカットの内容"):
            try:
                content = file_path_obj.read_text(encoding="utf-8")
                # URL=... の行を探す
                for line in content.splitlines():
                    if line.startswith("URL="):
                        target_url = line.split("URL=")[1].strip()
                        st.markdown(f"[このリンク先を開く]({target_url})")
                        st.text(f"URL: {target_url}")
                        break
            except Exception as e:
                st.error(f"URLファイルの読み込みエラー: {e}")
        return

    # その他ファイル → ダウンロード
    if file_path_obj.is_file():
        with st.expander(f"ファイルをダウンロード: {file_name}"):
            try:
                with file_path_obj.open("rb") as f:
                    st.download_button(
                        label="ダウンロード",
                        data=f,
                        file_name=file_name,
                        key=f"dl_{unique_key}"
                    )
            except Exception as e:
                st.warning(f"読み込みエラー: {e}")

# ==========================================
# エクスポート用ヘルパー
# ==========================================
def export_to_excel(entry):
    output = io.BytesIO()
    # 縦長形式で1つのレコードを書き出す
    details = entry.get("details", {})
    data = {
        "項目": ["タイトル", "登録日時", "場所", "参加者", "議題", "内容", "備考", "次回予定", "URL", "タグ", "添付ファイル"],
        "内容": [
            entry.get("title", ""),
            entry.get("created_at", ""),
            details.get("場所", ""),
            details.get("参加者", ""),
            details.get("議題", ""),
            details.get("内容", ""),
            details.get("備考", ""),
            details.get("次回予定", ""),
            entry.get("url", ""),
            ", ".join(entry.get("tags", [])),
            "\n".join(entry.get("file_path", [])) if isinstance(entry.get("file_path"), list) else entry.get("file_path", "")
        ]
    }
    df = pd.DataFrame(data)
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='単一議事録')
    return output.getvalue()

def export_all_to_excel(results):
    output = io.BytesIO()
    # 横長形式で全件を書き出す
    rows = []
    for entry in results:
        details = entry.get("details", {})
        rows.append({
            "タイトル": entry.get("title", ""),
            "登録日時": entry.get("created_at", ""),
            "日付": entry.get("date", ""),
            "時間": entry.get("time", ""),
            "場所": details.get("場所", ""),
            "参加者": details.get("参加者", ""),
            "議題": details.get("議題", ""),
            "内容": details.get("内容", ""),
            "備考": details.get("備考", ""),
            "次回予定": details.get("次回予定", ""),
            "URL": entry.get("url", ""),
            "タグ": ", ".join(entry.get("tags", [])),
            "添付ファイル": "\n".join(entry.get("file_path", [])) if isinstance(entry.get("file_path"), list) else entry.get("file_path", "")
        })
    df = pd.DataFrame(rows)
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='検索結果一覧')
    return output.getvalue()

def export_to_word(entry):
    output = io.BytesIO()
    doc = Document()
    
    # 余白の設定（1.5cmに縮小して1枚に収まりやすくする）
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # 全体のフォント設定 (10.5pt)
    style = doc.styles['Normal']
    style.font.name = '游明朝' # 日本語フォント指定（環境に合わせて調整）
    style.font.size = Pt(10.5)
    # 段落間隔を詰める
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    # タイトル
    doc.add_heading(entry.get("title", "議事録"), 0)
    
    # 基本情報テーブル
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    # 列幅の設定（左側を狭く、右側を広く）
    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(5.0)
    
    details = entry.get("details", {})
    info_list = [
        ("日時", details.get("日時", entry.get("created_at", ""))),
        ("場所", details.get("場所", "")),
        ("参加者", details.get("参加者", "")),
        ("次回予定", details.get("次回予定", "")),
        ("参考URL", entry.get("url", ""))
    ]
    
    for label, value in info_list:
        if value:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            # 各セルの幅を明示的に指定
            row_cells[0].width = Inches(1.0)
            row_cells[1].width = Inches(5.0)
            
    doc.add_paragraph()
    
    # 議題 (見出しサイズを小さく: Level 3)
    if details.get("議題"):
        h_agenda = doc.add_heading("議題", level=2)
        h_agenda.style.font.size = Pt(12)
        doc.add_paragraph(details.get("議題"))
        
    # 内容
    h_content = doc.add_heading("内容", level=2)
    h_content.style.font.size = Pt(12)
    content_text = details.get("内容", entry.get("text", ""))
    doc.add_paragraph(content_text)
    
    doc.save(output)
    return output.getvalue()

def export_to_pdf(entry):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # 日本語フォントの設定 (Windows の標準フォントを使用)
    font_path = r"C:\Windows\Fonts\msmincho.ttc"
    # MS Mincho は TTC なので、index 指定が必要な場合があるが fpdf2 はパスで概ね通る
    # 通らない場合はフォント名を指定して追加
    pdf.add_font("MSMincho", "", font_path)
    pdf.set_font("MSMincho", size=10.5)

    # タイトル
    pdf.set_font("MSMincho", size=16)
    pdf.cell(0, 10, entry.get("title", "議事録"), ln=True, align='C')
    pdf.ln(5)

    # 基本情報テーブル
    pdf.set_font("MSMincho", size=10.5)
    details = entry.get("details", {})
    info_list = [
        ("日時", details.get("日時", entry.get("created_at", ""))),
        ("場所", details.get("場所", "")),
        ("参加者", details.get("参加者", "")),
        ("次回予定", details.get("次回予定", "")),
        ("参考URL", entry.get("url", ""))
    ]

    for label, value in info_list:
        if value:
            # ラベル列 (30mm), 内容列 (160mm)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(30, 8, label, border=1, fill=True)
            pdf.cell(160, 8, str(value), border=1, ln=True)

    pdf.ln(5)

    # 議題
    if details.get("議題"):
        pdf.set_font("MSMincho", size=12)
        pdf.cell(0, 10, "議題", ln=True)
        pdf.set_font("MSMincho", size=10.5)
        pdf.multi_cell(0, 8, details.get("議題"), border=0)
        pdf.ln(3)

    # 内容
    pdf.set_font("MSMincho", size=12)
    pdf.cell(0, 10, "内容", ln=True)
    pdf.set_font("MSMincho", size=10.5)
    content_text = details.get("内容", entry.get("text", ""))
    pdf.multi_cell(0, 8, content_text, border=0)

    return bytes(pdf.output())

def export_to_zip(entry):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        # 1. PDF を生成して追加
        try:
            pdf_data = export_to_pdf(entry)
            pdf_name = f"議事録_{entry.get('title', '無題')}.pdf"
            zf.writestr(pdf_name, pdf_data)
        except Exception as e:
            # PDF失敗時はテキストファイルで代替
            zf.writestr("error_log.txt", f"PDF生成エラー: {e}")

        # 2. 添付ファイルを追加
        raw_file_paths = entry.get("file_path", [])
        if isinstance(raw_file_paths, str):
            file_list = [raw_file_paths] if raw_file_paths.strip() else []
        else:
            file_list = raw_file_paths if isinstance(raw_file_paths, list) else []

        for fp in file_list:
            path_obj = Path(fp.strip())
            if path_obj.exists() and path_obj.is_file():
                # ZIP 内ではファイル名のみにする
                zf.write(path_obj, arcname=path_obj.name)
            
    return zip_buffer.getvalue()

# ==========================================
# 結果表示（編集・削除機能付き）
# ==========================================
def display_results(results, data_dir):
    existing_tags = collect_existing_tags()

    for i, entry in enumerate(results):
        record_id = entry.get("id")
        source_type = entry.get("source", "mobile")
        unique_key = f"{record_id}_{i}"

        st.markdown("----")

        # 情報源のバッジ表示
        source_labels = {"mobile": "📱通常記録", "minutes": "📝議事録", "reference": "📚参考データ"}
        st.caption(f"情報源: {source_labels.get(source_type, source_type)}")

        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"### {entry.get('title', '無題')}")
        with col2:
            st.text(f"登録日時: {entry.get('created_at', '')}")

        st.markdown(f"**内容:**\n{entry.get('text', '')}")

        tags = entry.get("tags", [])
        if tags:
            st.markdown(f"**タグ:** {', '.join(tags)}")

        # URLフィールドの表示
        url = entry.get("url", "")
        if url:
            st.markdown(f"**参考リンク:** [{url}]({url})")

        # ファイルプレビュー (複数ファイル対応)
        raw_file_paths = entry.get("file_path", "")
        if isinstance(raw_file_paths, str):
            file_list = [raw_file_paths] if raw_file_paths.strip() else []
        else:
            file_list = raw_file_paths if isinstance(raw_file_paths, list) else []

        if file_list:
            if len(file_list) > 1:
                st.info(f"添付ファイル: {len(file_list)}件")
            for j, fp in enumerate(file_list):
                show_file_preview(fp.strip(), f"{unique_key}_{j}")

        # ==========================================
        # 編集・削除・出力ボタン
        # ==========================================
        if record_id is not None:
            btn_col1, btn_col2, btn_col3, btn_col4, btn_col5, btn_col6 = st.columns([1, 1, 2, 2, 2, 2])

            with btn_col1:
                edit_clicked = st.button("編集", key=f"edit_{unique_key}")
            with btn_col2:
                delete_clicked = st.button("削除", key=f"delete_{unique_key}")
            
            with btn_col3:
                # Word出力
                word_data = export_to_word(entry)
                st.download_button(
                    label="Word出力",
                    data=word_data,
                    file_name=f"議事録_{entry.get('title', '無題')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"word_{unique_key}"
                )
            
            with btn_col4:
                # PDF出力
                try:
                    pdf_data = export_to_pdf(entry)
                    st.download_button(
                        label="PDF出力",
                        data=pdf_data,
                        file_name=f"議事録_{entry.get('title', '無題')}.pdf",
                        mime="application/pdf",
                        key=f"pdf_{unique_key}"
                    )
                except Exception as e:
                    st.error(f"PDF生成エラー: {e}")

            with btn_col5:
                # Excel出力
                excel_data = export_to_excel(entry)
                st.download_button(
                    label="Excel出力",
                    data=excel_data,
                    file_name=f"議事録_{entry.get('title', '無題')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key=f"excel_{unique_key}"
                )

            with btn_col6:
                # ZIP出力 (配布用)
                try:
                    zip_data = export_to_zip(entry)
                    st.download_button(
                        label="配布用ZIP",
                        data=zip_data,
                        file_name=f"議事録セット_{entry.get('title', '無題')}.zip",
                        mime="application/zip",
                        key=f"zip_{unique_key}"
                    )
                except Exception as e:
                    st.error(f"ZIP生成エラー: {e}")

            # --- 編集モード ---
            edit_state_key = f"editing_{unique_key}"
            if edit_clicked:
                st.session_state[edit_state_key] = True

            if st.session_state.get(edit_state_key, False):
                with st.container():
                    st.markdown("---")
                    st.subheader("記録の編集")

                    new_title = st.text_input("タイトル", value=entry.get("title", ""), key=f"ed_title_{unique_key}")
                    new_url = st.text_input("URL", value=entry.get("url", ""), key=f"ed_url_{unique_key}")
                    
                    details = entry.get("details", {})
                    new_details = {}
                    
                    if details:
                        st.markdown("**詳細項目の編集**")
                        # 日時のパースと編集
                        dt_str = details.get("日時", "")
                        d_val = datetime.now().date()
                        ts_val = datetime.now().time()
                        te_val = datetime.now().time()
                        
                        # YYYY-MM-DD HH:MM ～ HH:MM の形式を試行
                        range_match = re.search(r"(\d{4}-\d{2}-\d{2}) (\d{2}:\d{2}) ～ (\d{2}:\d{2})", dt_str)
                        if range_match:
                            try:
                                d_val = datetime.strptime(range_match.group(1), "%Y-%m-%d").date()
                                ts_val = datetime.strptime(range_match.group(2), "%H:%M").time()
                                te_val = datetime.strptime(range_match.group(3), "%H:%M").time()
                            except: pass
                        else:
                            # YYYY-MM-DD HH:MM の形式を試行
                            single_match = re.search(r"(\d{4}-\d{2}-\d{2}) (\d{2}:\d{2})", dt_str)
                            if single_match:
                                try:
                                    d_val = datetime.strptime(single_match.group(1), "%Y-%m-%d").date()
                                    ts_val = datetime.strptime(single_match.group(2), "%H:%M").time()
                                    te_val = ts_val # 終了時刻がなければ開始と同じ
                                except: pass

                        col_ed_d, col_ed_ts, col_ed_te = st.columns([2, 1, 1])
                        with col_ed_d:
                            ed_d = st.date_input("日付", value=d_val, key=f"ed_date_{unique_key}")
                        with col_ed_ts:
                            ed_ts = st.time_input("開始", value=ts_val, key=f"ed_ts_{unique_key}")
                        with col_ed_te:
                            ed_te = st.time_input("終了", value=te_val, key=f"ed_te_{unique_key}")
                        new_details["日時"] = f"{ed_d} {ed_ts.strftime('%H:%M')} ～ {ed_te.strftime('%H:%M')}"

                        # 参加者の編集
                        st.markdown("参加者")
                        existing_p = collect_existing_participants(data_dir)
                        current_p_str = details.get("参加者", "")
                        current_p_list = [p.strip() for p in re.split(r'[,\n、，\s]+', current_p_str) if p.strip()]
                        
                        # 既存の選択肢に現在の参加者も追加
                        p_options = sorted(set(existing_p) | set(current_p_list))
                        ed_p_multi = st.multiselect("参加者を選択", options=p_options, default=current_p_list, key=f"ed_p_multi_{unique_key}")
                        ed_p_new = st.text_input("参加者を追加（カンマ区切り）", key=f"ed_p_new_{unique_key}")
                        
                        p_added = [p.strip() for p in re.split(r'[，、, ]+', ed_p_new.replace("　", ",")) if p.strip()]
                        final_p_list = list(dict.fromkeys(ed_p_multi + p_added))
                        new_details["参加者"] = ", ".join(final_p_list)

                        # その他のフィールド（固定項目を優先）
                        other_fields = ["場所", "議題", "内容", "備考", "次回予定"]
                        # 実際にdetailsにあるキーも含める
                        all_detail_keys = list(dict.fromkeys(other_fields + list(details.keys())))
                        
                        for k in all_detail_keys:
                            if k in ["日時", "参加者"]: continue
                            val = details.get(k, "")
                            if len(str(val)) > 50 or k in ["議題", "内容", "備考"]:
                                new_details[k] = st.text_area(k, value=str(val), key=f"ed_det_{k}_{unique_key}")
                            else:
                                new_details[k] = st.text_input(k, value=str(val), key=f"ed_det_{k}_{unique_key}")

                        # プレビュー用の全文（text）を再生成するか、元のtextを編集するか
                        # ここでは、detailsがある場合はdetailsから再生成するのがminutes.pyの仕様に合う
                        content_lines = []
                        for k, v in new_details.items():
                            if str(v).strip():
                                content_lines.append(f"【{k}】\n{v}\n")
                        new_text_val = "\n".join(content_lines)
                        new_text = st.text_area("内容（プレビュー・自由編集）", value=new_text_val, key=f"ed_text_{unique_key}", height=300)
                    else:
                        new_text = st.text_area("内容", value=entry.get("text", ""), key=f"ed_text_{unique_key}", height=300)

                    new_files = st.file_uploader("ファイルを追加", type=["jpg", "jpeg", "png", "pdf", "csv", "xlsx", "txt", "bat", "url", "mp4", "mov", "avi", "mkv"], accept_multiple_files=True, key=f"ed_files_{unique_key}")
                    enable_video_compression_ed = st.checkbox("動画を最適化（リサイズ・圧縮）する", value=True, help="動画のサイズを縮小し、再生互換性を高めます", key=f"ed_compress_{unique_key}")

                    # タグ編集
                    current_tags = entry.get("tags", [])
                    all_tag_options = sorted(set(existing_tags) | set(current_tags))
                    edited_tags = st.multiselect("タグを選択", options=all_tag_options, default=current_tags, key=f"ed_tags_{unique_key}")
                    extra_tags_input = st.text_input("新しいタグを追加（カンマ区切り）", key=f"ed_newtags_{unique_key}")

                    save_col, cancel_col, _ = st.columns([1, 1, 6])
                    with save_col:
                        if st.button("保存", key=f"save_{unique_key}"):
                            # 新規タグ処理
                            extra_processed = (
                                extra_tags_input.strip()
                                .replace("　", ",").replace(" ", ",")
                                .replace("、", ",").replace("，", ",")
                            )
                            extra_list = [t.strip() for t in extra_processed.split(",") if t.strip()]
                            final_tags = list(dict.fromkeys(edited_tags + extra_list))

                            # 追加ファイル用パス
                            current_file_paths = entry.get("file_path", [])
                            if isinstance(current_file_paths, str):
                                current_file_paths = [current_file_paths] if current_file_paths else []
                            
                            new_file_paths = []
                            if new_files:
                                image_dir = data_dir / "images"
                                image_dir.mkdir(parents=True, exist_ok=True)
                                
                                for j, uploaded_file in enumerate(new_files):
                                    timestamp_str = datetime.now().strftime("%Y%m%d%H%M%S")
                                    ext = Path(uploaded_file.name).suffix
                                    safe_title = new_title.replace("\\", "_").replace("/", "_").replace(":", "_").replace("*", "_")
                                    safe_title = safe_title.replace("?", "_").replace("\"", "_").replace("<", "_").replace(">", "_").replace("|", "_")
                                    safe_title = safe_title.replace(" ", "_").replace("　", "_").replace("、", "_").replace("，", "_")
                                    filename = f"{timestamp_str}_{safe_title}_edit_{j}{ext}"
                                    file_save_path = image_dir / filename
                                    
                                    try:
                                        # 動画圧縮の判定
                                        is_video = ext.lower() in [".mp4", ".mov", ".avi", ".mkv"]
                                        if is_video and enable_video_compression_ed:
                                            with st.spinner(f"動画を圧縮中: {uploaded_file.name}..."):
                                                temp_path = image_dir / f"temp_{filename}"
                                                with temp_path.open("wb") as f:
                                                    f.write(uploaded_file.getbuffer())
                                                
                                                if ext.lower() != ".mp4":
                                                    filename = filename.rsplit(".", 1)[0] + ".mp4"
                                                    file_save_path = image_dir / filename

                                                success, err_msg = compress_video(temp_path, file_save_path)
                                                if temp_path.exists(): temp_path.unlink()
                                                
                                                if not success:
                                                    st.error(f"動画の圧縮に失敗しました: {err_msg}")
                                                    with file_save_path.open("wb") as f:
                                                        f.write(uploaded_file.getbuffer())
                                        else:
                                            with file_save_path.open("wb") as f:
                                                f.write(uploaded_file.getbuffer())
                                        
                                        new_file_paths.append(str(file_save_path).replace("\\", "/"))
                                    except Exception as e:
                                        st.error(f"ファイルの保存に失敗しました ({uploaded_file.name}): {e}")

                            updated = {
                                "title": new_title,
                                "url": new_url,
                                "text": new_text,
                                "details": new_details if new_details else entry.get("details", {}),
                                "tags": final_tags,
                                "file_path": current_file_paths + new_file_paths  # 既存リストに結合
                                    }
                            if update_entry(record_id, updated):
                                st.success("更新しました！")
                                st.session_state[edit_state_key] = False
                                # 検索結果をクリアして再検索を促す
                                if "results" in st.session_state:
                                    del st.session_state["results"]
                                st.rerun()
                            else:
                                st.error("更新に失敗しました。")

                    with cancel_col:
                        if st.button("キャンセル", key=f"cancel_{unique_key}"):
                            st.session_state[edit_state_key] = False
                            st.rerun()

            # --- 削除確認 ---
            delete_confirm_key = f"del_confirm_{unique_key}"
            if delete_clicked:
                st.session_state[delete_confirm_key] = True

            if st.session_state.get(delete_confirm_key, False):
                st.warning(f"「{entry.get('title', '無題')}」を本当に削除しますか？（※データベースからのみ削除され、添付ファイルは保持されます）")
                yes_col, no_col, _ = st.columns([1, 1, 6])
                with yes_col:
                    if st.button("はい、削除する", key=f"yes_del_{unique_key}"):
                        if delete_entry(record_id):
                            st.success("削除しました。")
                            st.session_state[delete_confirm_key] = False
                            if "results" in st.session_state:
                                del st.session_state["results"]
                            st.rerun()
                        else:
                            st.error("削除に失敗しました。")
                with no_col:
                    if st.button("キャンセル", key=f"no_del_{unique_key}"):
                        st.session_state[delete_confirm_key] = False
                        st.rerun()

# ==========================================
# 検索画面 UI
# ==========================================
def show_search():
    st.header("記録の検索")

    data_dir_path = st.session_state.get("data_dir", DEFAULT_DATA_DIR)
    data_dir = Path(data_dir_path)

    if not data_dir.exists():
        st.warning(f"データフォルダが存在しません: {data_dir}")
        return

    # タグサジェスト用のタグ一覧を取得
    existing_tags = collect_existing_tags(data_dir)

    with st.form("search_form"):
        col_s1, col_s2 = st.columns([2, 1])
        with col_s1:
            title = st.text_input("タイトル")
        with col_s2:
            source = st.selectbox("情報の種類", ["すべて", "mobile", "minutes", "reference"], format_func=lambda x: {"すべて":"すべて", "mobile":"通常記録", "minutes":"議事録", "reference":"参考データ"}.get(x, x))
        
        text = st.text_input("内容")
        tag = st.text_input("タグ")
        mode = st.radio("検索モード", ["or", "and"], horizontal=True)
        submitted = st.form_submit_button("検索")

    if submitted:
        conditions = {
            "title": title.strip(),
            "text": text.strip(),
            "tag": tag.strip(),
            "source": source
        }
        st.session_state["results"] = search_entries(conditions, mode)

    if "results" in st.session_state:
        results = st.session_state["results"]
        st.success(f"{len(results)} 件の記録が見つかりました")

        if results:
            # 一括エクスポート
            st.markdown("### 検索結果をまとめて保存")
            bulk_excel_data = export_all_to_excel(results)
            st.download_button(
                label="検索結果をエクセルで一括ダウンロード",
                data=bulk_excel_data,
                file_name="議事録検索結果一覧.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="bulk_excel"
            )

        display_results(results, data_dir)
